# Authors: Michal Yechezkel (ID: 322556267), Almog Talker (ID: 322546680)
"""Render the report's block model to PDF (ReportLab) and to DOCX (python-docx).

The report body lives in ``src/report_content.py`` as a list of plain dictionaries --
headings, paragraphs, tables, figures, call-out boxes -- so the same source produces both
deliverables and the two can never say different things. Inline emphasis uses a small
HTML-like subset (``<b>``, ``<i>``, ``<sub>``, ``<sup>``, ``<link>``, plus the handful of
named entities we need); ``_runs`` parses it once for both back-ends.
"""

from __future__ import annotations

import html
import os
import re
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

# ---------------------------------------------------------------------------
# Shared inline markup
# ---------------------------------------------------------------------------

ENTITIES = {
    "&plusmn;": "±", "&minus;": "−", "&times;": "×",
    "&mdash;": "—", "&ndash;": "–", "&nbsp;": " ",
    "&middot;": "·", "&deg;": "°", "&radic;": "√",
    "&sigma;": "σ", "&gamma;": "γ", "&nu;": "ν",
    "&Delta;": "Δ", "&rho;": "ρ",
    # NB: no "&ell;" -- U+2113 has no glyph in the base-14 Times font and renders as a
    # black box. Write the L1 norm as "<i>L</i><sub>1</sub>" instead.
    "&ldquo;": "“", "&rdquo;": "”", "&lsquo;": "‘",
    "&rsquo;": "’", "&hellip;": "…", "&amp;": "&",
    "&lt;": "<", "&gt;": ">", "&asymp;": "≈", "&le;": "≤",
    "&ge;": "≥", "&rarr;": "→", "&equiv;": "≡",
}

_TAG = re.compile(r"</?(?:b|i|sub|sup|u|tt|br\s*/?|link[^>]*)>", re.I)


def to_unicode(markup: str) -> str:
    """Replace our named entities with real characters, keeping the inline tags.

    ReportLab's paragraph parser knows only a subset of named entities (capital Greek
    letters, for instance, are silently dropped), so we normalise them ourselves and hand
    it plain UTF-8. ``&amp;`` is substituted last so it cannot re-create an entity.
    """
    text = markup
    for entity, char in ENTITIES.items():
        if entity in ("&amp;", "&lt;", "&gt;"):
            continue
        text = text.replace(entity, char)
    # ReportLab has no <tt>; unknown tags are dropped *with their content*, which silently
    # emptied table headers. Map it onto the monospace font tag it does understand.
    text = re.sub(r"<tt>", '<font name="Courier">', text, flags=re.I)
    text = re.sub(r"</tt>", "</font>", text, flags=re.I)
    return text


def plain_text(markup: str) -> str:
    """Strip the inline markup, keeping the characters (used for DOCX and for checks)."""
    text = _TAG.sub("", markup)
    for entity, char in ENTITIES.items():
        text = text.replace(entity, char)
    return text


def _runs(markup: str) -> List[Tuple[str, Dict[str, bool]]]:
    """Split inline markup into ``(text, {bold, italic, sub, sup})`` runs."""
    out: List[Tuple[str, Dict[str, bool]]] = []
    state = {"bold": False, "italic": False, "sub": False, "sup": False,
             "underline": False, "mono": False}
    token = re.compile(r"(</?(?:b|i|sub|sup|u|tt|br\s*/?|link[^>]*)>)", re.I)
    for piece in token.split(markup):
        if not piece:
            continue
        low = piece.lower()
        if low.startswith("<"):
            closing = low.startswith("</")
            name = low.lstrip("</").rstrip(">").split()[0].rstrip("/")
            key = {"b": "bold", "i": "italic", "sub": "sub", "sup": "sup",
                   "u": "underline", "tt": "mono"}.get(name)
            if name.startswith("br"):
                out.append(("\n", dict(state)))
            elif key:
                state[key] = not closing
            continue
        text = piece
        for entity, char in ENTITIES.items():
            text = text.replace(entity, char)
        text = html.unescape(text)
        out.append((text, dict(state)))
    return out


# ---------------------------------------------------------------------------
# PDF back-end
# ---------------------------------------------------------------------------


def render_pdf(blocks: Sequence[Dict], out_path: str) -> str:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (
        BaseDocTemplate, Frame, Image, KeepTogether, PageBreak, PageTemplate,
        Paragraph, Spacer, Table, TableStyle,
    )

    page_w, page_h = A4
    margin = 1.85 * cm
    content_w = page_w - 2 * margin

    ink = colors.HexColor("#1a1a1a")
    accent = colors.HexColor("#1f4e79")
    rule = colors.HexColor("#b9c6d4")
    soft = colors.HexColor("#f3f6fa")
    callout_bg = colors.HexColor("#fdf7e6")
    callout_edge = colors.HexColor("#d9c89a")

    base = getSampleStyleSheet()
    S = {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Times-Bold",
                                fontSize=19.5, leading=23, textColor=accent, spaceAfter=2),
        "subtitle": ParagraphStyle("subtitle", parent=base["Title"],
                                   fontName="Times-Roman", fontSize=13, leading=16.5,
                                   textColor=ink, spaceAfter=8),
        "meta": ParagraphStyle("meta", parent=base["Normal"], fontName="Times-Roman",
                               fontSize=10.2, leading=14, alignment=TA_CENTER,
                               textColor=ink),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="Times-Bold",
                             fontSize=14.2, leading=17.5, textColor=accent,
                             spaceBefore=14, spaceAfter=4),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Times-Bold",
                             fontSize=11.4, leading=14.6, textColor=ink,
                             spaceBefore=10, spaceAfter=3),
        "h3": ParagraphStyle("h3", parent=base["Heading3"], fontName="Times-BoldItalic",
                             fontSize=10.4, leading=13.6, textColor=ink,
                             spaceBefore=8, spaceAfter=2),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Times-Roman",
                               fontSize=10.05, leading=14.1, alignment=TA_JUSTIFY,
                               textColor=ink, spaceAfter=5),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"], fontName="Times-Roman",
                                 fontSize=10.05, leading=13.8, alignment=TA_JUSTIFY,
                                 textColor=ink, leftIndent=14, spaceAfter=3),
        "caption": ParagraphStyle("caption", parent=base["Normal"],
                                  fontName="Times-Roman", fontSize=8.65, leading=11.3,
                                  alignment=TA_JUSTIFY,
                                  textColor=colors.HexColor("#31363c"),
                                  spaceBefore=3, spaceAfter=9),
        "cell": ParagraphStyle("cell", parent=base["Normal"], fontName="Times-Roman",
                               fontSize=8.25, leading=10.3, textColor=ink),
        "cellhead": ParagraphStyle("cellhead", parent=base["Normal"],
                                   fontName="Times-Bold", fontSize=8.25, leading=10.3,
                                   textColor=colors.white),
        "callout": ParagraphStyle("callout", parent=base["Normal"],
                                  fontName="Times-Roman", fontSize=9.6, leading=13.2,
                                  alignment=TA_JUSTIFY, textColor=ink, spaceAfter=3),
    }

    def furniture(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(rule)
        canvas.setLineWidth(0.5)
        canvas.line(margin, margin - 0.3 * cm, page_w - margin, margin - 0.3 * cm)
        canvas.setFont("Times-Roman", 8.4)
        canvas.setFillColor(colors.HexColor("#59636e"))
        canvas.drawString(margin, margin - 0.8 * cm,
                          "Magnetic Resonance Imaging 361.2.6501 — Final Project")
        canvas.drawRightString(page_w - margin, margin - 0.8 * cm, str(doc.page))
        canvas.restoreState()

    doc = BaseDocTemplate(
        out_path, pagesize=A4, leftMargin=margin, rightMargin=margin, topMargin=margin,
        bottomMargin=margin + 0.55 * cm,
        title="MRI Restoration from Subsampled k-space — Final Project",
        author="Michal Yechezkel, Almog Talker",
        subject="Magnetic Resonance Imaging 361.2.6501")
    frame = Frame(margin, doc.bottomMargin, content_w,
                  page_h - margin - doc.bottomMargin, id="body")
    doc.addPageTemplates([PageTemplate(id="single", frames=[frame], onPage=furniture)])

    def para(markup: str, style: str):
        return Paragraph(to_unicode(markup), S[style])

    def hrule():
        t = Table([[""]], colWidths=[content_w], rowHeights=[1.05])
        t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), rule)] +
                              [(k, (0, 0), (-1, -1), 0) for k in
                               ("LEFTPADDING", "RIGHTPADDING", "TOPPADDING",
                                "BOTTOMPADDING")]))
        return t

    story: List = []
    for b in blocks:
        kind = b["type"]
        if kind == "title":
            story += [Spacer(1, 4), para(b["course"], "meta"),
                      para("<b>Final Project</b>", "meta"), Spacer(1, 9),
                      para(b["heading"], "title"), para(b["subheading"], "subtitle"),
                      para(b["authors"], "meta"), Spacer(1, 3),
                      para(b["link"], "meta"), Spacer(1, 11), hrule(), Spacer(1, 9)]
        elif kind == "h1":
            # keepWithNext stops a heading being orphaned at the foot of a page: ReportLab
            # pulls it onto the next page along with whatever follows it.
            group = [para(b["text"], "h1"), hrule(), Spacer(1, 5)]
            for flowable in group:
                flowable.keepWithNext = 1
            story += group
        elif kind in ("h2", "h3"):
            heading = para(b["text"], kind)
            heading.keepWithNext = 1
            story.append(heading)
        elif kind == "p":
            story.append(para(b["text"], "body"))
        elif kind == "bullets":
            story += [Paragraph(to_unicode("•&nbsp;&nbsp;" + t), S["bullet"])
                      for t in b["items"]]
            story.append(Spacer(1, 3))
        elif kind == "callout":
            inner = Table([[para("<b>" + b["title"] + "</b>", "callout")],
                           [para(b["body"], "callout")]],
                          colWidths=[content_w - 0.6 * cm])
            inner.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), callout_bg),
                ("BOX", (0, 0), (-1, -1), 0.7, callout_edge),
                ("LINEBEFORE", (0, 0), (0, -1), 2.6, accent),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, 0), 6),
                ("TOPPADDING", (0, 1), (-1, -1), 0),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 7),
            ]))
            story += [Spacer(1, 3), inner, Spacer(1, 7)]
        elif kind == "table":
            rows = b["rows"]
            n_cols = len(rows[0])
            widths = b.get("widths")
            if widths:
                total = sum(widths)
                col_widths = [content_w * w / total for w in widths]
            else:
                first = content_w * 0.26
                col_widths = [first] + [(content_w - first) / max(n_cols - 1, 1)] * (n_cols - 1)
            data = [[para(str(c), "cellhead" if r == 0 else "cell") for c in row]
                    for r, row in enumerate(rows)]
            t = Table(data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
            style = [
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("BOX", (0, 0), (-1, -1), 0.7, accent),
                ("LINEBELOW", (0, 0), (-1, 0), 0.7, accent),
                ("LINEBELOW", (0, 1), (-1, -2), 0.25, rule),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (b.get("align_right_from", 1), 1), (-1, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 2.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
            ]
            for i in range(2, len(rows), 2):
                style.append(("BACKGROUND", (0, i), (-1, i), soft))
            for i in b.get("highlight", ()):
                style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#e2efdd")))
            t.setStyle(TableStyle(style))
            group: List = [t]
            if b.get("caption"):
                group += [Spacer(1, 2), para(b["caption"], "caption")]
            story.append(KeepTogether(group))
        elif kind == "figure":
            path = b["path"]
            if not os.path.exists(path):
                story.append(para(f"<i>[missing figure: {path}]</i>", "caption"))
                continue
            iw, ih = ImageReader(path).getSize()
            width = content_w * b.get("width_frac", 1.0)
            height = width * ih / iw
            cap = b.get("max_height_cm", 20.0) * cm
            if height > cap:
                height = cap
                width = height * iw / ih
            img = Image(path, width=width, height=height)
            img.hAlign = "CENTER"
            group = [img]
            if b.get("caption"):
                group += [Spacer(1, 2), para(b["caption"], "caption")]
            story.append(KeepTogether(group) if b.get("keep", True) else group[0])
            if not b.get("keep", True):
                story += group[1:]
        elif kind == "pagebreak":
            story.append(PageBreak())
        elif kind == "spacer":
            story.append(Spacer(1, b.get("height", 6)))
        else:
            raise ValueError(f"unknown block type {kind!r}")

    doc.build(story)
    return out_path


# ---------------------------------------------------------------------------
# DOCX back-end
# ---------------------------------------------------------------------------


def render_docx(blocks: Sequence[Dict], out_path: str) -> str:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    ACCENT = RGBColor(0x1F, 0x4E, 0x79)
    INK = RGBColor(0x1A, 0x1A, 0x1A)
    CAPTION_INK = RGBColor(0x31, 0x36, 0x3C)

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.9)
    section.left_margin = section.right_margin = Cm(1.9)
    usable_cm = (section.page_width - section.left_margin - section.right_margin) / Cm(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    def shade(cell, hex_colour: str) -> None:
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hex_colour)
        cell._tc.get_or_add_tcPr().append(el)

    def add(markup: str, *, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            bold_all: bool = False, italic_all: bool = False,
            colour: Optional[RGBColor] = None, space_before: float = 0,
            space_after: float = 5, indent: float = 0, keep_with_next: bool = False):
        p = document.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.keep_with_next = keep_with_next
        if indent:
            pf.left_indent = Cm(indent)
        for text, flags in _runs(markup):
            if text == "\n":
                run = p.add_run()
                run.add_break()
                continue
            run = p.add_run(text)
            run.font.name = "Courier New" if flags["mono"] else "Times New Roman"
            run.font.size = Pt(size * (0.72 if (flags["sub"] or flags["sup"])
                                       else (0.92 if flags["mono"] else 1.0)))
            run.bold = flags["bold"] or bold_all
            run.italic = flags["italic"] or italic_all
            run.underline = flags["underline"]
            run.font.subscript = flags["sub"]
            run.font.superscript = flags["sup"]
            run.font.color.rgb = colour if colour is not None else INK
        return p

    def horizontal_rule():
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(5)
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "B9C6D4")
        borders.append(bottom)
        pPr.append(borders)

    # Running footer with the page number.
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run()
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(8.5)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    for b in blocks:
        kind = b["type"]
        if kind == "title":
            add(b["course"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
            add("<b>Final Project</b>", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
            add(b["heading"], size=19, bold_all=True, colour=ACCENT,
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            add(b["subheading"], size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=9)
            add(b["authors"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            add(b["link"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=9)
            horizontal_rule()
        elif kind == "h1":
            add(b["text"], size=14.5, bold_all=True, colour=ACCENT,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=13, space_after=2,
                keep_with_next=True)
            horizontal_rule()
        elif kind == "h2":
            add(b["text"], size=11.5, bold_all=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=10, space_after=3, keep_with_next=True)
        elif kind == "h3":
            add(b["text"], size=10.5, bold_all=True, italic_all=True,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=2,
                keep_with_next=True)
        elif kind == "p":
            add(b["text"])
        elif kind == "bullets":
            for item in b["items"]:
                add("• " + item, indent=0.5, space_after=3)
        elif kind == "callout":
            table = document.add_table(rows=2, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Cm(usable_cm - 0.4)
            for row_idx, markup, bold in ((0, b["title"], True), (1, b["body"], False)):
                cell = table.cell(row_idx, 0)
                cell.width = Cm(usable_cm - 0.4)
                shade(cell, "FDF7E6")
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for text, flags in _runs(markup):
                    run = para.add_run(text)
                    run.font.name = "Courier New" if flags["mono"] else "Times New Roman"
                    run.font.size = Pt(9.8 * (0.72 if (flags["sub"] or flags["sup"])
                                              else (0.92 if flags["mono"] else 1.0)))
                    run.bold = flags["bold"] or bold
                    run.italic = flags["italic"]
                    run.font.subscript = flags["sub"]
                    run.font.superscript = flags["sup"]
            add("", space_after=2)
        elif kind == "table":
            rows = b["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            widths = b.get("widths")
            if widths:
                total = sum(widths)
                col_cm = [usable_cm * w / total for w in widths]
            else:
                first = usable_cm * 0.26
                col_cm = [first] + [(usable_cm - first) / max(len(rows[0]) - 1, 1)] * (len(rows[0]) - 1)
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    cell = table.cell(r, c)
                    cell.width = Cm(col_cm[c])
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.space_before = Pt(1)
                    if c >= b.get("align_right_from", 1) and r > 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for text, flags in _runs(str(value)):
                        run = para.add_run(text)
                        run.font.name = "Courier New" if flags["mono"] else "Times New Roman"
                        run.font.size = Pt(8.4 * (0.75 if (flags["sub"] or flags["sup"])
                                                  else (0.92 if flags["mono"] else 1.0)))
                        run.bold = flags["bold"] or r == 0
                        run.italic = flags["italic"]
                        run.font.subscript = flags["sub"]
                        run.font.superscript = flags["sup"]
                        if r == 0:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    if r == 0:
                        shade(cell, "1F4E79")
                    elif r % 2 == 0:
                        shade(cell, "F3F6FA")
                    for i in b.get("highlight", ()):
                        if r == i:
                            shade(cell, "E2EFDD")
            if b.get("caption"):
                add(b["caption"], size=8.7, italic_all=False, colour=CAPTION_INK,
                    space_before=3, space_after=9)
        elif kind == "figure":
            path = b["path"]
            if not os.path.exists(path):
                add(f"<i>[missing figure: {path}]</i>", size=8.7, colour=CAPTION_INK)
                continue
            from PIL import Image as PILImage
            iw, ih = PILImage.open(path).size
            width_cm = usable_cm * b.get("width_frac", 1.0)
            height_cm = width_cm * ih / iw
            cap_cm = b.get("max_height_cm", 20.0)
            if height_cm > cap_cm:
                height_cm = cap_cm
                width_cm = height_cm * iw / ih
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(path, width=Cm(width_cm))
            if b.get("caption"):
                add(b["caption"], size=8.7, colour=CAPTION_INK, space_after=9)
        elif kind == "pagebreak":
            document.add_page_break()
        elif kind == "spacer":
            add("", space_after=b.get("height", 6))
        else:
            raise ValueError(f"unknown block type {kind!r}")

    document.save(out_path)
    return out_path
